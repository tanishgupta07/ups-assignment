from docx import Document as DocxDocument
from langchain_core.documents import Document

def extract_docx(path, metadata=None):
    doc = DocxDocument(path)
    content = []

    for element in doc.element.body:
        tag = element.tag.split('}')[-1]

        if tag == 'p':
            # paragraph
            para = element.text
            if para and para.strip():
                content.append(para)

        elif tag == 'tbl':
            # table - find it in doc.tables and extract as plain text
            for table in doc.tables:
                if table._tbl == element:
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        content.append(" | ".join(cells))
                    break

    text = "\n".join(content)
    return [Document(page_content=text, metadata=metadata or {})]
